#!/usr/bin/env python3
"""Build editable DOCX, verified PDF, supplement, and submission manifests."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def markdown_docx(source: Path, destination: Path, title: str) -> None:
    from docx import Document
    from docx.shared import Pt
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    text = source.read_text(encoding="utf-8-sig")
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if match:
            document.add_heading(match.group(2), level=min(len(match.group(1)), 3))
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(stripped)
    document.core_properties.title = title
    document.save(destination)


def convert_pdf(docx_path: Path, output_dir: Path, soffice: str) -> Path:
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        capture_output=True, text=True, encoding="utf-8", timeout=90,
    )
    pdf = output_dir / f"{docx_path.stem}.pdf"
    if result.returncode or not pdf.is_file() or not pdf.read_bytes().startswith(b"%PDF-"):
        raise RuntimeError(result.stderr or result.stdout or "PDF conversion failed")
    return pdf


def build(run_dir: Path, manuscript: Path, title: str, soffice: str) -> dict:
    output_dir = run_dir / "09_论文正文" / "投稿文件_publication_files"
    output_dir.mkdir(parents=True, exist_ok=True)
    manuscript_docx = output_dir / "论文稿_manuscript.docx"
    markdown_docx(manuscript, manuscript_docx, title)
    manuscript_pdf = convert_pdf(manuscript_docx, output_dir, soffice)

    supplement_md = output_dir / "补充材料_supplement.md"
    supplement_md.write_text(
        "# 补充材料\n\n## 数据处理与计分\n\n见数据派生报告及冻结清单。\n\n"
        "## 测量不变性\n\n报告configural、metric、scalar及必要的partial invariance。\n\n"
        "## 稳健性分析\n\n报告零值密集两部分模型、学校固定效应及替代计分规则。\n",
        encoding="utf-8", newline="\n",
    )
    supplement_docx = output_dir / "补充材料_supplement.docx"
    markdown_docx(supplement_md, supplement_docx, f"{title} - Supplement")
    supplement_pdf = convert_pdf(supplement_docx, output_dir, soffice)

    cover_letter = output_dir / "投稿信_cover_letter.md"
    cover_letter.write_text(
        f"# Cover Letter\n\nDear Editor,\n\nPlease consider the manuscript “{title}”.\n\n"
        "This package is a preparation draft. Journal name, novelty statement, ethics, funding, conflicts, and data availability must be verified before submission.\n",
        encoding="utf-8", newline="\n",
    )
    manifest_csv = output_dir / "表图清单_tables_figures_manifest.csv"
    candidates = [path for stage in [run_dir / "07_统计分析", run_dir / "08_结果与图表"] if stage.exists() for path in stage.rglob("*") if path.suffix.lower() in {".csv", ".png", ".svg", ".jpg", ".jpeg"}]
    with manifest_csv.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["file", "type", "sha256", "bytes", "caption_status"])
        writer.writeheader()
        for path in sorted(candidates):
            writer.writerow({"file": str(path.resolve()), "type": "figure" if path.suffix.lower() != ".csv" else "table", "sha256": sha256(path), "bytes": path.stat().st_size, "caption_status": "pending-verification"})
    files = [manuscript_docx, manuscript_pdf, supplement_md, supplement_docx, supplement_pdf, cover_letter, manifest_csv]
    manifest = {
        "schema_version": 1, "status": "complete", "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "files": [{"path": str(path.resolve()), "sha256": sha256(path), "bytes": path.stat().st_size} for path in files],
        "submission_status": "preparation-only-requires-journal-policy-and-claim-audit",
    }
    submission_manifest = output_dir / "投稿文件清单_submission_manifest.json"
    submission_manifest.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {
        "status": "complete", "manuscript_docx": str(manuscript_docx.resolve()),
        "manuscript_pdf": str(manuscript_pdf.resolve()), "supplement_docx": str(supplement_docx.resolve()),
        "supplement_pdf": str(supplement_pdf.resolve()), "cover_letter": str(cover_letter.resolve()),
        "tables_figures_manifest": str(manifest_csv.resolve()), "submission_manifest": str(submission_manifest.resolve()),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--run-dir", required=True)
    parser.add_argument("--manuscript", required=True)
    parser.add_argument("--title", required=True)
    parser.add_argument("--soffice", default=shutil.which("soffice") or "soffice")
    args = parser.parse_args()
    try:
        result = build(Path(args.run_dir).resolve(), Path(args.manuscript).resolve(), args.title, args.soffice)
    except (OSError, RuntimeError) as error:
        print(json.dumps({"status": "blocked", "error": str(error)}, ensure_ascii=False))
        return 3
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
